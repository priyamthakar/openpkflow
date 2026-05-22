"""Word document report renderers for openpkflow dissolution reports."""

from __future__ import annotations

import base64
import io
from datetime import datetime, timezone
from pathlib import Path
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from openpkflow.nca.results import NCAResult, NCASummaryResults

_DISCLAIMER = (
    "This report was generated using OpenPKFlow — an open-source Python workflow "
    "for pharmacometric analysis. Final regulatory interpretation should be reviewed "
    "by qualified formulation, pharmacokinetic, and regulatory experts."
)

_FIT_DISCLAIMER = (
    "Dissolution model fitting characterises the release mechanism of a formulation — "
    "it is not a regulatory similarity test. For dissolution similarity assessment, "
    "use the f2 or bootstrap f2 method (FDA 1997 guidance)."
)


def render_comparison_docx_report(
    *,
    title: str,
    reference_label: str,
    test_label: str,
    f1_value: float,
    f2_value: float,
    n_timepoints: int,
    time_points: list[float],
    reference_mean: list[float],
    test_mean: list[float],
    output_path: str | Path | None = None,
) -> bytes:
    """Render a dissolution comparison report as a Word document.

    Parameters
    ----------
    title :
        Report title.
    reference_label :
        Label for the reference formulation.
    test_label :
        Label for the test formulation.
    f1_value :
        Computed f1 difference factor.
    f2_value :
        Computed f2 similarity factor.
    n_timepoints :
        Number of matched timepoints used in the calculation.
    time_points :
        Timepoint values used in the comparison.
    reference_mean :
        Mean percent dissolved values for the reference at each timepoint.
    test_mean :
        Mean percent dissolved values for the test at each timepoint.
    output_path :
        If given, write the DOCX bytes to this path (parent dirs created
        automatically).

    Returns
    -------
    bytes
        DOCX file contents as bytes.
    """
    try:
        import docx  # noqa: F401
    except ImportError as exc:
        raise ImportError(
            "Word export requires python-docx. Install with: pip install openpkflow[reports]"
        ) from exc

    from docx import Document
    from docx.shared import Inches, Pt

    from openpkflow import __version__
    from openpkflow.dissolution.plotting import dissolution_profile_plot_b64

    if f2_value >= 50:
        interpretation = "f2 >= 50 supports similarity between the reference and test profiles."
    else:
        interpretation = "f2 < 50 does not support similarity."

    plot_b64 = dissolution_profile_plot_b64(
        time_points=time_points,
        reference_mean=reference_mean,
        test_mean=test_mean,
        reference_label=reference_label,
        test_label=test_label,
    )

    document = Document()

    document.add_heading(title, level=1)

    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta_para = document.add_paragraph(
        f"Generated: {generated_at}  |  OpenPKFlow v{__version__}"
    )
    meta_para.runs[0].font.size = Pt(9)

    document.add_heading("Summary", level=2)

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, heading in enumerate(["Parameter", "Value"]):
        hdr_cells[i].text = heading
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    summary_rows = [
        ("Reference", reference_label),
        ("Test", test_label),
        ("Timepoints", str(n_timepoints)),
        ("f1 (difference factor)", f"{f1_value:.2f}"),
        ("f2 (similarity factor)", f"{f2_value:.2f}"),
        ("Interpretation", interpretation),
    ]
    for label, value in summary_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    document.add_heading("Dissolution Profile", level=2)

    img_data = base64.b64decode(plot_b64)
    img_buf = io.BytesIO(img_data)
    document.add_picture(img_buf, width=Inches(5))

    document.add_heading("Data Table", level=2)

    data_table = document.add_table(rows=1, cols=3)
    data_table.style = "Table Grid"

    data_hdr_cells = data_table.rows[0].cells
    for i, heading in enumerate(["Time (min)", reference_label, test_label]):
        data_hdr_cells[i].text = heading
    for cell in data_hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for tp, ref, tst in zip(time_points, reference_mean, test_mean, strict=True):
        row_cells = data_table.add_row().cells
        row_cells[0].text = str(tp)
        row_cells[1].text = f"{ref:.2f}"
        row_cells[2].text = f"{tst:.2f}"

    document.add_paragraph()

    disclaimer_para = document.add_paragraph(_DISCLAIMER)
    for run in disclaimer_para.runs:
        run.italic = True
        run.font.size = Pt(9)

    buf = io.BytesIO()
    document.save(buf)
    docx_bytes = buf.getvalue()

    if output_path is not None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(docx_bytes)

    return docx_bytes


def render_multi_media_docx_report(
    *,
    title: str,
    reference_label: str,
    test_label: str,
    media_names: list[str],
    per_media_results: dict[str, dict[str, Any]],
    f2_summary: dict[str, float],
    overall_pass: bool,
    plot_b64: str,
    output_path: str | Path | None = None,
) -> bytes:
    """Render a multi-media dissolution report as a Word document.

    Parameters
    ----------
    title :
        Report title.
    reference_label :
        Label for the reference formulation.
    test_label :
        Label for the test formulation.
    media_names :
        Ordered list of medium names.
    per_media_results :
        Dict mapping medium name to its ComparisonResult.to_dict() output.
    f2_summary :
        Dict mapping medium name to its f2 value.
    overall_pass :
        True if all media have f2 >= 50.
    plot_b64 :
        Base64-encoded PNG of the multi-media panel plot.
    output_path :
        If given, write the DOCX bytes to this path.

    Returns
    -------
    bytes
        DOCX file contents as bytes.
    """
    try:
        import docx  # noqa: F401
    except ImportError as exc:
        raise ImportError(
            "Word export requires python-docx. Install with: pip install openpkflow[reports]"
        ) from exc

    from docx import Document
    from docx.shared import Inches, Pt

    from openpkflow import __version__

    document = Document()

    document.add_heading(title, level=1)

    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta_para = document.add_paragraph(
        f"Generated: {generated_at}  |  OpenPKFlow v{__version__}"
    )
    meta_para.runs[0].font.size = Pt(9)

    # Overall verdict
    verdict_para = document.add_paragraph()
    verdict_run = verdict_para.add_run(
        "Overall: PASS" if overall_pass else "Overall: FAIL"
    )
    verdict_run.bold = True
    verdict_run.font.size = Pt(12)

    document.add_heading("Multi-Media f2 Summary", level=2)

    summary_table = document.add_table(rows=1, cols=5)
    summary_table.style = "Table Grid"
    hdr = summary_table.rows[0].cells
    for i, heading in enumerate(["Medium", "f2", "f1", "Timepoints", "Status"]):
        hdr[i].text = heading
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    for medium in media_names:
        cr = per_media_results.get(medium, {})
        if not cr:
            continue
        f2 = cr.get("f2_value", 0)
        row = summary_table.add_row().cells
        row[0].text = medium
        row[1].text = f"{f2:.2f}"
        row[2].text = f"{cr.get('f1_value', 0):.2f}"
        row[3].text = str(cr.get("n_timepoints", 0))
        row[4].text = "PASS" if f2 >= 50 else "FAIL"

    document.add_paragraph()

    document.add_heading("Multi-Media Profile Plot", level=2)
    img_data = base64.b64decode(plot_b64)
    img_buf = io.BytesIO(img_data)
    document.add_picture(img_buf, width=Inches(6))

    document.add_paragraph()

    disclaimer_para = document.add_paragraph(_DISCLAIMER)
    for run in disclaimer_para.runs:
        run.italic = True
        run.font.size = Pt(9)

    buf = io.BytesIO()
    document.save(buf)
    docx_bytes = buf.getvalue()

    if output_path is not None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(docx_bytes)

    return docx_bytes


def render_model_fit_docx_report(
    *,
    formulation_label: str,
    time_points: list[float],
    observed_mean: list[float],
    fit_rows: list[dict[str, Any]],
    plot_b64: str,
    output_path: str | Path | None = None,
) -> bytes:
    """Render a dissolution model fit report as a Word document.

    Parameters
    ----------
    formulation_label :
        Label of the fitted formulation.
    time_points :
        Observed time points (minutes).
    observed_mean :
        Mean percent dissolved at each observed time point.
    fit_rows :
        Pre-processed fit data rows (one dict per model). Each dict must have:
        ``model_name``, ``params``, ``r_squared``, ``aic``, ``aicc``, ``bic``,
        ``n_points``, ``n_params``, ``converged``, ``rank``, ``is_best``.
    plot_b64 :
        Base64-encoded PNG of the model fit overlay plot.
    output_path :
        If given, write the DOCX bytes to this path (parent dirs created
        automatically).

    Returns
    -------
    bytes
        DOCX file contents as bytes.
    """
    try:
        import docx  # noqa: F401
    except ImportError as exc:
        raise ImportError(
            "Word export requires python-docx. Install with: pip install openpkflow[reports]"
        ) from exc

    from docx import Document
    from docx.shared import Inches, Pt

    from openpkflow import __version__

    def _fmt_params(params: dict[str, Any]) -> str:
        return ", ".join(f"{k}={v:.4g}" for k, v in params.items())

    converged_rows = sorted(
        [r for r in fit_rows if r.get("converged", False)],
        key=lambda r: r.get("rank", 9999),
    )
    failed_rows = [r for r in fit_rows if not r.get("converged", False)]

    title = f"Dissolution Model Fitting: {formulation_label}"
    document = Document()

    document.add_heading(title, level=1)

    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta_para = document.add_paragraph(
        f"Generated: {generated_at}  |  OpenPKFlow v{__version__}"
    )
    meta_para.runs[0].font.size = Pt(9)

    document.add_heading("Model Fit Results", level=2)

    if converged_rows:
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for i, heading in enumerate(["Rank", "Model", "R²", "AICc", "BIC", "Params"]):
            hdr_cells[i].text = heading
        for cell in hdr_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        for row in converged_rows:
            row_cells = table.add_row().cells
            row_cells[0].text = str(row.get("rank", ""))
            row_cells[1].text = str(row.get("model_name", ""))
            row_cells[2].text = f"{row.get('r_squared', float('nan')):.4f}"
            row_cells[3].text = f"{row.get('aicc', float('nan')):.2f}"
            row_cells[4].text = f"{row.get('bic', float('nan')):.2f}"
            row_cells[5].text = _fmt_params(row.get("params", {}))
            if row.get("is_best", False):
                for cell in row_cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
    else:
        document.add_paragraph("No models converged successfully.")

    if failed_rows:
        failed_names = ", ".join(r.get("model_name", "unknown") for r in failed_rows)
        failed_para = document.add_paragraph(f"Failed to converge: {failed_names}")
        failed_para.runs[0].font.size = Pt(9)
        failed_para.runs[0].italic = True

    document.add_heading("Fit Overlay Plot", level=2)

    img_data = base64.b64decode(plot_b64)
    img_buf = io.BytesIO(img_data)
    document.add_picture(img_buf, width=Inches(5))

    document.add_heading("Observed Data", level=2)

    data_table = document.add_table(rows=1, cols=2)
    data_table.style = "Table Grid"

    data_hdr_cells = data_table.rows[0].cells
    for i, heading in enumerate(["Time (min)", "Mean % Dissolved"]):
        data_hdr_cells[i].text = heading
    for cell in data_hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for tp, obs in zip(time_points, observed_mean, strict=True):
        row_cells = data_table.add_row().cells
        row_cells[0].text = str(tp)
        row_cells[1].text = f"{obs:.2f}"

    document.add_paragraph()

    disclaimer_para = document.add_paragraph(_DISCLAIMER)
    for run in disclaimer_para.runs:
        run.italic = True
        run.font.size = Pt(9)

    fit_disclaimer_para = document.add_paragraph(_FIT_DISCLAIMER)
    for run in fit_disclaimer_para.runs:
        run.italic = True
        run.font.size = Pt(9)

    buf = io.BytesIO()
    document.save(buf)
    docx_bytes = buf.getvalue()

    if output_path is not None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(docx_bytes)

    return docx_bytes


def render_nca_single_docx_report(
    *,
    result: NCAResult,
    output_path: str | Path | None = None,
) -> bytes:
    """Render a per-subject NCA report as a Word document.

    Parameters
    ----------
    result :
        NCA result for a single subject.
    output_path :
        If given, write the DOCX bytes to this path (parent dirs created automatically).

    Returns
    -------
    bytes
        DOCX file contents as bytes.
    """
    try:
        import docx  # noqa: F401
    except ImportError as exc:
        raise ImportError(
            "Word export requires python-docx. Install with: pip install openpkflow[reports]"
        ) from exc

    from docx import Document
    from docx.shared import Pt

    from openpkflow import __version__

    def _fmt(v: float | None) -> str:
        return f"{v:.4g}" if v is not None else "N/A"

    def _cl_label(r: NCAResult) -> tuple[str, str]:
        if r.CL_F is not None:
            return "CL_F (L/h)", _fmt(r.CL_F)
        if r.CL is not None:
            return "CL (L/h)", _fmt(r.CL)
        return "CL_F (L/h)", "N/A"

    def _vz_label(r: NCAResult) -> tuple[str, str]:
        if r.Vz_F is not None:
            return "Vz_F (L)", _fmt(r.Vz_F)
        if r.Vz is not None:
            return "Vz (L)", _fmt(r.Vz)
        return "Vz_F (L)", "N/A"

    cl_lbl, cl_val = _cl_label(result)
    vz_lbl, vz_val = _vz_label(result)
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    document = Document()
    document.add_heading(f"NCA Report - Subject {result.subject}", level=1)

    meta_para = document.add_paragraph(
        f"Generated: {generated_at}  |  OpenPKFlow v{__version__}"
    )
    meta_para.runs[0].font.size = Pt(9)

    document.add_heading("Study Parameters", level=2)
    study_table = document.add_table(rows=1, cols=2)
    study_table.style = "Table Grid"
    hdr = study_table.rows[0].cells
    for i, h in enumerate(["Parameter", "Value"]):
        hdr[i].text = h
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for label, value in [
        ("Subject", str(result.subject)),
        ("Route", result.route),
        ("Dose", f"{result.dose:.4g}"),
        ("AUC Method", result.auc_method),
        ("BLQ Method", result.blq_method),
    ]:
        row_cells = study_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    document.add_heading("PK Parameters", level=2)
    pk_table = document.add_table(rows=1, cols=2)
    pk_table.style = "Table Grid"
    hdr2 = pk_table.rows[0].cells
    for i, h in enumerate(["Parameter", "Value"]):
        hdr2[i].text = h
    for cell in hdr2:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for label, value in [
        ("Cmax", _fmt(result.Cmax)),
        ("Tmax", _fmt(result.Tmax)),
        ("AUClast", _fmt(result.AUClast)),
        ("AUCinf_obs", _fmt(result.AUCinf_obs)),
        ("AUC % Extrapolated", _fmt(result.AUC_percent_extrapolated)),
        ("lambda_z", _fmt(result.lambda_z)),
        ("Half-life", _fmt(result.half_life)),
        ("lambda_z method", result.lambda_z_method or "N/A"),
        (cl_lbl, cl_val),
        (vz_lbl, vz_val),
    ]:
        row_cells = pk_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    if result.warnings:
        document.add_heading("Warnings", level=2)
        for w in result.warnings:
            document.add_paragraph(w, style="List Bullet")

    document.add_paragraph()
    disclaimer_para = document.add_paragraph(_DISCLAIMER)
    for run in disclaimer_para.runs:
        run.italic = True
        run.font.size = Pt(9)

    buf = io.BytesIO()
    document.save(buf)
    docx_bytes = buf.getvalue()

    if output_path is not None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(docx_bytes)

    return docx_bytes


def render_nca_summary_docx_report(
    *,
    summary: NCASummaryResults,
    output_path: str | Path | None = None,
) -> bytes:
    """Render a multi-subject NCA summary report as a Word document.

    Parameters
    ----------
    summary :
        Collection of per-subject NCA results.
    output_path :
        If given, write the DOCX bytes to this path (parent dirs created automatically).

    Returns
    -------
    bytes
        DOCX file contents as bytes.
    """
    try:
        import docx  # noqa: F401
    except ImportError as exc:
        raise ImportError(
            "Word export requires python-docx. Install with: pip install openpkflow[reports]"
        ) from exc

    from docx import Document
    from docx.shared import Pt

    from openpkflow import __version__

    def _fmt(v: float | None) -> str:
        return f"{v:.4g}" if v is not None else "N/A"

    def _cl_val(r: NCAResult) -> str:
        if r.CL_F is not None:
            return _fmt(r.CL_F)
        if r.CL is not None:
            return _fmt(r.CL)
        return "N/A"

    def _vz_val(r: NCAResult) -> str:
        if r.Vz_F is not None:
            return _fmt(r.Vz_F)
        if r.Vz is not None:
            return _fmt(r.Vz)
        return "N/A"

    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    document = Document()
    document.add_heading("NCA Summary Report", level=1)

    meta_para = document.add_paragraph(
        f"Generated: {generated_at}  |  OpenPKFlow v{__version__}"
    )
    meta_para.runs[0].font.size = Pt(9)

    if summary.study_label:
        document.add_heading("Study Parameters", level=2)
        info_table = document.add_table(rows=1, cols=2)
        info_table.style = "Table Grid"
        hdr = info_table.rows[0].cells
        for i, h in enumerate(["Parameter", "Value"]):
            hdr[i].text = h
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        for label, value in [
            ("Study", summary.study_label),
            ("AUC Method", summary.auc_method),
            ("BLQ Method", summary.blq_method),
        ]:
            row_cells = info_table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

    document.add_heading("PK Parameters by Subject", level=2)

    col_headers = [
        "Subject", "AUClast", "AUCinf_obs", "AUC%Extr",
        "Cmax", "Tmax", "Half-life", "CL/CL_F", "Vz/Vz_F",
    ]
    tbl = document.add_table(rows=1, cols=len(col_headers))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(col_headers):
        hdr_cells[i].text = h
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(8)

    for r in summary.results:
        row_cells = tbl.add_row().cells
        for i, val in enumerate([
            str(r.subject),
            _fmt(r.AUClast),
            _fmt(r.AUCinf_obs),
            _fmt(r.AUC_percent_extrapolated),
            _fmt(r.Cmax),
            _fmt(r.Tmax),
            _fmt(r.half_life),
            _cl_val(r),
            _vz_val(r),
        ]):
            row_cells[i].text = val
            for para in row_cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    document.add_paragraph()
    disclaimer_para = document.add_paragraph(_DISCLAIMER)
    for run in disclaimer_para.runs:
        run.italic = True
        run.font.size = Pt(9)

    buf = io.BytesIO()
    document.save(buf)
    docx_bytes = buf.getvalue()

    if output_path is not None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(docx_bytes)

    return docx_bytes


# ---------------------------------------------------------------------------
# Simulation report
# ---------------------------------------------------------------------------


def render_sim_docx_report(
    *,
    result: Any,
    output_path: str | Path | None = None,
    time_unit: str = "h",
    conc_unit: str = "ng/mL",
) -> bytes:
    """Render a PK simulation Word document report.

    Parameters
    ----------
    result : SimulationResult
        Simulation result to render.
    output_path : str | Path or None, optional
        If given, write the DOCX to this path.
    time_unit : str, optional
        Time unit label for tables and headings.
    conc_unit : str, optional
        Concentration unit label for tables and headings.

    Returns
    -------
    bytes
        DOCX byte content.
    """
    from datetime import datetime, timezone

    from docx import Document
    from docx.shared import Pt, RGBColor

    from openpkflow import __version__
    from openpkflow.sim.plotting import pk_profile_plot_b64
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    document = Document()
    _NAVY_RGB = RGBColor(0x0D, 0x3B, 0x66)

    model_name = type(result.model).__name__
    label = result.label or "N/A"
    route = result.regimen.route
    n_doses = len(result.regimen.doses)

    title_para = document.add_heading(f"PK Simulation Report -- {label}", 0)
    for run in title_para.runs:
        run.font.color.rgb = _NAVY_RGB

    meta_para = document.add_paragraph(
        f"Generated: {generated_at} | OpenPKFlow v{__version__} | "
        f"Model: {model_name} | Route: {route} | Doses: {n_doses}"
    )
    for run in meta_para.runs:
        run.font.size = Pt(9)

    document.add_heading("Simulation Results", 1)
    result_tbl = document.add_table(rows=1, cols=3)
    result_tbl.style = "Table Grid"
    hdr = result_tbl.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    hdr[2].text = "Unit"
    for metric, val, unit in [
        ("Cmax", f"{result.Cmax:.4g}", conc_unit),
        ("Tmax", f"{result.Tmax:.4g}", time_unit),
        ("Cmin", f"{min(result.concs):.4g}", conc_unit),
        ("Clast", f"{result.concs[-1]:.4g}", conc_unit),
        ("Time range", f"{result.times[0]:.4g} - {result.times[-1]:.4g}", time_unit),
    ]:
        row = result_tbl.add_row().cells
        row[0].text = metric
        row[1].text = val
        row[2].text = unit

    document.add_heading("Concentration-Time Profile", 1)
    b64 = pk_profile_plot_b64(
        times=result.times, concs=result.concs,
        dose_times=result.regimen.dose_times, label=result.label,
        time_unit=time_unit, conc_unit=conc_unit,
    )
    img_bytes = base64.b64decode(b64)
    document.add_picture(io.BytesIO(img_bytes))

    document.add_heading("Model Parameters", 1)
    params = result.model.param_dict()
    param_tbl = document.add_table(rows=1, cols=2)
    param_tbl.style = "Table Grid"
    param_hdr = param_tbl.rows[0].cells
    param_hdr[0].text = "Parameter"
    param_hdr[1].text = "Value"
    for k, v in params.items():
        row = param_tbl.add_row().cells
        row[0].text = str(k)
        row[1].text = f"{v:.4g}" if isinstance(v, float) else str(v)

    document.add_heading(f"Dose Regimen ({n_doses} dose{'s' if n_doses != 1 else ''})", 1)
    dose_tbl = document.add_table(rows=1, cols=3)
    dose_tbl.style = "Table Grid"
    dose_hdr = dose_tbl.rows[0].cells
    dose_hdr[0].text = "#"
    dose_hdr[1].text = f"Time ({time_unit})"
    dose_hdr[2].text = "Amount"
    for i, d in enumerate(result.regimen.doses):
        t_inf_str = f" (inf {d.t_inf:.4g} h)" if d.t_inf is not None else ""
        row = dose_tbl.add_row().cells
        row[0].text = str(i + 1)
        row[1].text = f"{d.time:.4g}"
        row[2].text = f"{d.amount:.4g}{t_inf_str}"

    if result.warnings:
        document.add_heading("Warnings", 1)
        for w in result.warnings:
            document.add_paragraph(f"- {w}")

    document.add_paragraph()
    disclaimer_para = document.add_paragraph(_DISCLAIMER)
    for run in disclaimer_para.runs:
        run.italic = True
        run.font.size = Pt(9)

    report_buf = io.BytesIO()
    document.save(report_buf)
    docx_bytes = report_buf.getvalue()

    if output_path is not None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(docx_bytes)

    return docx_bytes


def render_ivivc_docx_report(
    *,
    result: object,
    output_path: str | Path | None = None,
) -> bytes:
    """Render an IVIVC Word document report.

    Parameters
    ----------
    result : IVIVCResult
        IVIVC result to render.
    output_path : str or Path or None, optional
        If given, write DOCX to this path.

    Returns
    -------
    bytes
        DOCX byte content.
    """
    try:
        import docx  # noqa: F401
    except ImportError as exc:
        raise ImportError(
            "Word export requires python-docx. Install with: pip install openpkflow[reports]"
        ) from exc

    from docx import Document
    from docx.shared import Inches, Pt

    import base64 as _b64
    import io as _io
    import matplotlib as _mpl
    _mpl.use("Agg")
    import matplotlib.pyplot as _plt
    import numpy as _np

    from openpkflow import __version__

    lp = result.levy_plot
    pp = result.predictability

    # Generate plot
    fig, axes = _plt.subplots(2, 2, figsize=(8, 6), dpi=200)
    ax = axes[0, 0]
    ax.plot(result.times, result.fa, "o-", color="#003366", linewidth=2, markersize=4)
    ax.plot(result.ivt_times, result.ivt_fraction, "s--", color="#cc3300", linewidth=1, markersize=3, label="In vitro")
    ax.set_title("Fraction Absorbed vs Dissolved", fontsize=9)
    ax.set_xlabel("Time (h)", fontsize=8)
    ax.set_ylabel("Fraction", fontsize=8)
    ax.legend(fontsize=7)
    ax.grid(True, alpha=0.3)
    ax.set_ylim(0, 1.1)
    ax = axes[0, 1]
    ax.scatter(lp["x"], lp["y"], color="#003366", s=20, zorder=3)
    if len(lp.get("x", [])) > 1:
        x_line = _np.linspace(0, 1, 100)
        y_line = lp["slope"] * x_line + lp["intercept"]
        ax.plot(x_line, y_line, "-", color="#cc3300", linewidth=1, label=f"R2={lp['r_squared']:.3f}")
        ax.plot([0, 1], [0, 1], ":", color="#888888", linewidth=1, label="1:1")
    ax.set_title("Levy Plot", fontsize=9)
    ax.set_xlabel("In vitro F_d", fontsize=8)
    ax.set_ylabel("In vivo F_a", fontsize=8)
    ax.legend(fontsize=6)
    ax.grid(True, alpha=0.3)
    ax.set_xlim(0, 1.05)
    ax.set_ylim(0, 1.05)
    ax = axes[1, 0]
    ax.plot(result.times, result.concentrations, "o-", color="#003366", linewidth=2, markersize=4, label="Observed")
    ax.plot(result.predicted_times, result.predicted_concs, "--", color="#cc3300", linewidth=1, label="Predicted")
    ax.set_title("Predicted vs Observed", fontsize=9)
    ax.set_xlabel("Time (h)", fontsize=8)
    ax.set_ylabel("Concentration", fontsize=8)
    ax.legend(fontsize=7)
    ax.grid(True, alpha=0.3)
    ax = axes[1, 1]
    ax.plot(result.ivt_times, result.ivt_fraction * 100, "o-", color="#006699", linewidth=2, markersize=4)
    ax.set_title("Dissolution Profile", fontsize=9)
    ax.set_xlabel("Time (min)", fontsize=8)
    ax.set_ylabel("% Dissolved", fontsize=8)
    ax.grid(True, alpha=0.3)
    ax.set_ylim(0, 105)
    fig.tight_layout()
    plot_buf = _io.BytesIO()
    fig.savefig(plot_buf, format="png", bbox_inches="tight", dpi=200)
    _plt.close(fig)
    plot_buf.seek(0)

    title = "IVIVC Level A Report"
    if result.study_label:
        title += f" -- {result.study_label}"
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    document = Document()
    document.add_heading(title, level=1)
    meta_para = document.add_paragraph(f"Generated: {generated_at}  |  OpenPKFlow v{__version__}")
    meta_para.runs[0].font.size = Pt(9)

    document.add_heading("Predictability Assessment (FDA 1997)", level=2)
    overall = "PASS" if pp.get("overall_pass", False) else "FAIL"
    pp_table = document.add_table(rows=1, cols=4)
    pp_table.style = "Table Grid"
    for i, h in enumerate(["Metric", "Value", "Criterion", "Status"]):
        pp_table.rows[0].cells[i].text = h
        for p in pp_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for label, value, crit, status in [
        ("Cmax %PE", f"{pp.get('%PE_Cmax', 0):.2f}%", "<= 15%", "PASS" if pp.get("passes_cmax", False) else "FAIL"),
        ("AUCinf %PE", f"{pp.get('%PE_AUC', 0):.2f}%", "<= 15%", "PASS" if pp.get("passes_auc", False) else "FAIL"),
        ("Mean abs %PE", f"{pp.get('mean_abs_%PE', 0):.2f}%", "<= 10%", "PASS" if pp.get("passes_mean", False) else "FAIL"),
        ("Overall", "", "", overall),
    ]:
        row = pp_table.add_row().cells
        row[0].text = label
        row[1].text = value
        row[2].text = crit
        row[3].text = status

    document.add_heading("IVIVC Plots", level=2)
    document.add_picture(plot_buf, width=Inches(5.5))
    document.add_heading("Levy Plot Regression", level=2)
    levy_table = document.add_table(rows=1, cols=2)
    levy_table.style = "Table Grid"
    for i, h in enumerate(["Metric", "Value"]):
        levy_table.rows[0].cells[i].text = h
    for label, value in [
        ("Slope", f"{lp.get('slope', 0):.4f}"),
        ("Intercept", f"{lp.get('intercept', 0):.4f}"),
        ("R-squared", f"{lp.get('r_squared', 0):.4f}"),
        ("N (0.05-0.95)", str(len(lp.get("x", [])))),
    ]:
        row = levy_table.add_row().cells
        row[0].text = label
        row[1].text = value

    document.add_paragraph()
    disclaimer_para = document.add_paragraph(_DISCLAIMER)
    for run in disclaimer_para.runs:
        run.italic = True
        run.font.size = Pt(9)

    report_buf = _io.BytesIO()
    document.save(report_buf)
    docx_bytes = report_buf.getvalue()
    if output_path is not None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(docx_bytes)
    return docx_bytes


def render_gof_docx_report(
    *,
    result: object,
    output_path: str | Path | None = None,
) -> bytes:
    """Render a population PK GOF Word document report.

    Parameters
    ----------
    result : GOFResult
        Computed GOF result.
    output_path : str | Path or None, optional
        Where to save the DOCX. If None, returns bytes without saving.

    Returns
    -------
    bytes
        DOCX bytes.
    """

    from docx import Document
    from docx.shared import Pt, RGBColor

    from openpkflow.pop.plotting import gof_plots_b64

    document = Document()
    _NAVY_RGB = RGBColor(0x0D, 0x3B, 0x66)

    title = "Population PK GOF Report"
    if result.study_label:
        title += f" -- {result.study_label}"
    title_para = document.add_heading(title, 0)
    for run in title_para.runs:
        run.font.color.rgb = _NAVY_RGB
    document.add_paragraph(
        f"Generated by OpenPKFlow | "
        f"{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )

    document.add_heading("GOF Metrics", 1)
    pm = result.pred_metrics()
    im = result.ipred_metrics()
    m_tbl = document.add_table(rows=1, cols=3)
    m_tbl.style = "Table Grid"
    hdr = m_tbl.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "OBS vs PRED"
    hdr[2].text = "OBS vs IPRED"
    for metric, pv, iv in [
        ("N", str(int(pm["n"])), str(int(im["n"]))),
        ("MPE (bias)", f"{pm['MPE']:.4g}", f"{im['MPE']:.4g}"),
        ("RMSE", f"{pm['RMSE']:.4g}", f"{im['RMSE']:.4g}"),
        ("rRMSE (%)", f"{pm['rRMSE_pct']:.2f}", f"{im['rRMSE_pct']:.2f}"),
        ("R2", f"{pm['R2']:.4f}", f"{im['R2']:.4f}"),
    ]:
        row = m_tbl.add_row().cells
        row[0].text = metric
        row[1].text = pv
        row[2].text = iv

    document.add_heading("GOF Diagnostic Plots", 1)
    b64 = gof_plots_b64(result)
    img_bytes = base64.b64decode(b64)
    document.add_picture(io.BytesIO(img_bytes))

    if result.warnings:
        document.add_heading("Warnings", 1)
        for w in result.warnings:
            document.add_paragraph(f"- {w}")

    document.add_paragraph()
    dis_para = document.add_paragraph(_DISCLAIMER)
    for run in dis_para.runs:
        run.italic = True
        run.font.size = Pt(9)

    report_buf = io.BytesIO()
    document.save(report_buf)
    docx_bytes = report_buf.getvalue()

    if output_path is not None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(docx_bytes)

    return docx_bytes


def render_vpc_docx_report(
    *,
    result: object,
    output_path: str | Path | None = None,
) -> bytes:
    """Render a VPC Word document report.

    Parameters
    ----------
    result : VPCResult
        Computed VPC result.
    output_path : str | Path or None, optional
        Where to save the DOCX. If None, returns bytes without saving.

    Returns
    -------
    bytes
        DOCX bytes.
    """
    import math

    from docx import Document
    from docx.shared import Pt, RGBColor

    from openpkflow.pop.plotting import vpc_plot_b64

    document = Document()
    _NAVY_RGB_VPC = RGBColor(0x0D, 0x3B, 0x66)

    title = "Visual Predictive Check (VPC)"
    if result.study_label:
        title += f" -- {result.study_label}"
    title_para = document.add_heading(title, 0)
    for run in title_para.runs:
        run.font.color.rgb = _NAVY_RGB_VPC
    document.add_paragraph(
        f"Generated by OpenPKFlow | "
        f"{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )

    document.add_heading("VPC Settings", 1)
    pi = result.pi
    s_tbl = document.add_table(rows=1, cols=2)
    s_tbl.style = "Table Grid"
    s_hdr = s_tbl.rows[0].cells
    s_hdr[0].text = "Setting"
    s_hdr[1].text = "Value"
    for k, v in [
        ("N observed", str(len(result.obs_times))),
        ("N replicates", str(result.n_replicates)),
        ("N bins", str(result.n_bins)),
        ("Percentiles", f"{pi[0]:.0f} / {pi[1]:.0f} / {pi[2]:.0f}"),
    ]:
        row = s_tbl.add_row().cells
        row[0].text = k
        row[1].text = v

    document.add_heading("Visual Predictive Check", 1)
    b64 = vpc_plot_b64(result)
    img_bytes = base64.b64decode(b64)
    document.add_picture(io.BytesIO(img_bytes))

    document.add_heading("VPC Band Data", 1)

    def _fv(v: float) -> str:
        return f"{v:.3g}" if not math.isnan(v) else "---"

    b_tbl = document.add_table(rows=1, cols=7)
    b_tbl.style = "Table Grid"
    b_hdr = b_tbl.rows[0].cells
    for i, h in enumerate([
        "Bin Mid",
        f"Obs {pi[0]:.0f}th", f"Obs {pi[1]:.0f}th", f"Obs {pi[2]:.0f}th",
        f"Sim {pi[0]:.0f}th", f"Sim {pi[1]:.0f}th", f"Sim {pi[2]:.0f}th",
    ]):
        b_hdr[i].text = h

    for i, mid in enumerate(result.bin_mids):
        row = b_tbl.add_row().cells
        row[0].text = f"{mid:.2f}"
        row[1].text = _fv(result.obs_lower[i])
        row[2].text = _fv(result.obs_median[i])
        row[3].text = _fv(result.obs_upper[i])
        row[4].text = _fv(result.sim_lower[i])
        row[5].text = _fv(result.sim_median[i])
        row[6].text = _fv(result.sim_upper[i])

    if result.warnings:
        document.add_heading("Warnings", 1)
        for w in result.warnings:
            document.add_paragraph(f"- {w}")

    document.add_paragraph()
    dis_para = document.add_paragraph(_DISCLAIMER)
    for run in dis_para.runs:
        run.italic = True
        run.font.size = Pt(9)

    report_buf = io.BytesIO()
    document.save(report_buf)
    docx_bytes = report_buf.getvalue()

    if output_path is not None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(docx_bytes)

    return docx_bytes
