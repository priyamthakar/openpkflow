"""Reports and reproducibility bundles for the dissolution workbench."""

from __future__ import annotations

import csv
import hashlib
import html
import io
import json
import zipfile
from pathlib import Path
from typing import TYPE_CHECKING, Literal

from openpkflow import __version__
from openpkflow.dissolution.workbench import workbench_result_json

if TYPE_CHECKING:
    from openpkflow.dissolution.workbench import DissolutionWorkbenchResult

_DISCLAIMER = (
    "This report was generated using OpenPKFlow (open-source). Final regulatory "
    "interpretation should be reviewed by qualified formulation, pharmacokinetic, "
    "and regulatory experts."
)


def _profile_plot_png(result: DissolutionWorkbenchResult) -> bytes:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    figure, axis = plt.subplots(figsize=(8.2, 4.8), dpi=160)
    for profile in result.reference_vessels:
        axis.plot(
            profile.time_points,
            profile.percent_released,
            color="#2c6eaa",
            alpha=0.22,
            linewidth=1.0,
        )
    for profile in result.test_vessels:
        axis.plot(
            profile.time_points,
            profile.percent_released,
            color="#d76a3a",
            alpha=0.22,
            linewidth=1.0,
        )
    axis.plot(
        result.comparison.time_points,
        result.comparison.reference_mean,
        "o-",
        color="#154c79",
        linewidth=2.4,
        label=f"{result.config.reference_label} mean",
    )
    axis.plot(
        result.comparison.time_points,
        result.comparison.test_mean,
        "s--",
        color="#b6401b",
        linewidth=2.4,
        label=f"{result.config.test_label} mean",
    )
    axis.axhline(85.0, color="#6b7280", linestyle=":", linewidth=1.0, label="85% threshold")
    axis.set_xlabel("Time (min)")
    axis.set_ylabel("Percent dissolved")
    axis.set_ylim(0.0, 105.0)
    axis.grid(alpha=0.22)
    axis.legend(loc="lower right", fontsize=8)
    figure.tight_layout()
    stream = io.BytesIO()
    figure.savefig(stream, format="png", bbox_inches="tight")
    plt.close(figure)
    return stream.getvalue()


def _model_rows(result: DissolutionWorkbenchResult) -> list[list[str]]:
    rows: list[list[str]] = []
    for label, fits in (
        (result.config.reference_label, result.reference_models.fits),
        (result.config.test_label, result.test_models.fits),
    ):
        ranked = sorted(
            (fit for fit in fits if fit.converged),
            key=lambda fit: fit.aicc,
        )
        for rank, fit in enumerate(ranked, 1):
            params = ", ".join(f"{key}={value:.5g}" for key, value in fit.params.items())
            rows.append(
                [
                    label,
                    str(rank),
                    fit.model_name,
                    params,
                    f"{fit.r_squared:.5f}",
                    f"{fit.aicc:.3f}",
                ]
            )
    return rows


def render_workbench_html(result: DissolutionWorkbenchResult) -> str:
    """Render a complete self-contained workbench HTML report.

    Parameters
    ----------
    result : DissolutionWorkbenchResult
        Completed workbench result.

    Returns
    -------
    str
        Self-contained HTML document.
    """
    import base64

    payload = result.to_dict()
    plot = base64.b64encode(_profile_plot_png(result)).decode("ascii")
    vessel_rows = "\n".join(
        "<tr>"
        f"<td>{html.escape(str(row['formulation']))}</td>"
        f"<td>{html.escape(str(row['batch']))}</td>"
        f"<td>{float(row['time']):g}</td>"
        f"<td>{float(row['percent_released']):.3f}</td>"
        "</tr>"
        for row in result.normalized_rows
    )
    model_rows = "\n".join(
        "<tr>" + "".join(f"<td>{html.escape(value)}</td>" for value in row) + "</tr>"
        for row in _model_rows(result)
    )
    warning_items = "".join(f"<li>{html.escape(item)}</li>" for item in result.warnings)
    if not warning_items:
        warning_items = "<li>None captured.</li>"
    config_json = html.escape(json.dumps(payload["config"], indent=2, sort_keys=True))
    comparison = result.comparison
    bootstrap = result.bootstrap
    model_comparison = result.model_comparison
    point_decision = (
        "Supports similarity" if comparison.f2_value >= 50 else "Does not support similarity"
    )
    bootstrap_decision = (
        "Supports similarity" if bootstrap.is_similar else "Does not support similarity"
    )
    msd_decision = (
        "Supports similarity" if result.msd_result.is_similar else "Does not support similarity"
    )
    model_decision = (
        "Supports similarity" if model_comparison.is_similar else "Does not support similarity"
    )
    return f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Advanced Dissolution Workbench Report</title>
<style>
:root {{ color-scheme: light; --navy:#123b5d; --blue:#dcebf7; --line:#d5dee7; }}
body {{ margin:0; background:#edf2f7; color:#1f2937; font:14px/1.5 Arial,sans-serif; }}
.page {{ max-width:1050px; margin:24px auto; background:white; box-shadow:0 8px 28px #1f293722; }}
header {{ background:var(--navy); color:white; padding:30px 38px; }}
header h1 {{ margin:0 0 8px; font-size:27px; }}
header p {{ margin:0; opacity:.88; }}
main {{ padding:30px 38px 40px; }}
h2 {{ color:var(--navy); margin-top:30px; border-bottom:2px solid var(--blue);
      padding-bottom:6px; }}
.metrics {{ display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); gap:12px; }}
.metric {{ border:1px solid var(--line); border-radius:7px; padding:14px; background:#f8fbfd; }}
.metric b {{ display:block; color:var(--navy); font-size:20px; }}
table {{ width:100%; border-collapse:collapse; margin:12px 0 20px; }}
th,td {{ border:1px solid var(--line); padding:7px 9px; text-align:left; }}
th {{ background:var(--blue); color:var(--navy); }}
pre {{ overflow:auto; background:#f5f7f9; border:1px solid var(--line); padding:14px; }}
img {{ display:block; max-width:100%; margin:12px auto; }}
.disclaimer {{ margin-top:30px; border-left:4px solid var(--navy); padding:12px 16px;
               background:#f5f8fb; }}
@media(max-width:720px) {{
  .page {{ margin:0; }} main,header {{ padding:22px 18px; }}
  .metrics {{ grid-template-columns:repeat(2,minmax(0,1fr)); }}
  table {{ display:block; overflow-x:auto; }}
}}
</style>
</head>
<body><div class="page">
<header>
  <h1>Advanced Dissolution Workbench</h1>
  <p>{html.escape(result.config.reference_label)} vs
     {html.escape(result.config.test_label)} | OpenPKFlow {__version__}</p>
</header>
<main>
<p><strong>Generated (UTC):</strong> {html.escape(result.generated_at_utc)}</p>
<h2>Decision Summary</h2>
<div class="metrics">
  <div class="metric">f1<b>{comparison.f1_value:.2f}</b></div>
  <div class="metric">f2 ({html.escape(comparison.f2_method)})<b>{comparison.f2_value:.2f}</b></div>
  <div class="metric">Bootstrap {bootstrap.confidence_level:.0%} CI
    <b>{bootstrap.ci_lower:.2f} - {bootstrap.ci_upper:.2f}</b></div>
  <div class="metric">Max deviation<b>{result.maximum_deviation:.2f}</b></div>
</div>
<table>
<thead><tr><th>Method</th><th>Result</th><th>Decision</th></tr></thead>
<tbody>
<tr><td>Point f2</td><td>{comparison.f2_value:.3f}</td>
<td>{point_decision}</td></tr>
<tr><td>Bootstrap f2 (all points)</td><td>{bootstrap.ci_lower:.3f} to
{bootstrap.ci_upper:.3f}</td>
<td>{bootstrap_decision}</td></tr>
<tr><td>Mahalanobis statistical distance</td>
<td>{result.msd_result.msd_squared:.3f} (critical {result.msd_result.chi2_05_critical:.3f})</td>
<td>{msd_decision}</td></tr>
<tr><td>Model-dependent {html.escape(model_comparison.param_name)}</td>
<td>{model_comparison.ratio_pct:.2f}% (90% CI {model_comparison.ci_lo:.2f} -
{model_comparison.ci_hi:.2f}%)</td>
<td>{model_decision}</td></tr>
</tbody></table>

<h2>Mean and Vessel Profiles</h2>
<img alt="Dissolution vessel and mean profile plot" src="data:image/png;base64,{plot}">

<h2>Five-Model AICc Ranking</h2>
<table><thead><tr><th>Formulation</th><th>Rank</th><th>Model</th>
<th>Parameters</th><th>R2</th><th>AICc</th></tr></thead>
<tbody>{model_rows}</tbody></table>

<h2>Normalized Vessel-Level Input</h2>
<table><thead><tr><th>Formulation</th><th>Vessel</th><th>Time</th>
<th>Percent released</th></tr></thead><tbody>{vessel_rows}</tbody></table>

<h2>Warnings and Prerequisites</h2>
<ul>{warning_items}</ul>

<h2>Exact Configuration</h2>
<pre>{config_json}</pre>

<div class="disclaimer"><strong>Disclaimer.</strong> {html.escape(_DISCLAIMER)}</div>
</main></div></body></html>"""


def _write_pdf(result: DissolutionWorkbenchResult, output_path: Path) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    styles = getSampleStyleSheet()
    story: list[object] = [
        Paragraph("Advanced Dissolution Workbench", styles["Title"]),
        Paragraph(
            f"{html.escape(result.config.reference_label)} vs "
            f"{html.escape(result.config.test_label)} | OpenPKFlow {__version__}",
            styles["Normal"],
        ),
        Spacer(1, 10),
        Table(
            [
                ["Metric", "Value"],
                ["f1", f"{result.comparison.f1_value:.3f}"],
                [
                    f"f2 ({result.comparison.f2_method})",
                    f"{result.comparison.f2_value:.3f}",
                ],
                [
                    f"Bootstrap {result.bootstrap.confidence_level:.0%} CI",
                    f"{result.bootstrap.ci_lower:.3f} - {result.bootstrap.ci_upper:.3f}",
                ],
                ["Maximum deviation", f"{result.maximum_deviation:.3f}"],
                [
                    "MSD squared / critical",
                    f"{result.msd_result.msd_squared:.3f} / "
                    f"{result.msd_result.chi2_05_critical:.3f}",
                ],
            ],
            colWidths=[2.7 * inch, 2.7 * inch],
        ),
        Spacer(1, 12),
        Image(io.BytesIO(_profile_plot_png(result)), width=6.7 * inch, height=3.9 * inch),
        Spacer(1, 12),
        Paragraph("Five-Model AICc Ranking", styles["Heading2"]),
    ]
    model_table = Table(
        [["Formulation", "Rank", "Model", "Parameters", "R2", "AICc"], *_model_rows(result)],
        repeatRows=1,
        colWidths=[1.0 * inch, 0.4 * inch, 1.0 * inch, 1.9 * inch, 0.65 * inch, 0.65 * inch],
    )
    model_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dcebf7")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#aebdcc")),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.extend(
        [
            model_table,
            PageBreak(),
            Paragraph("Normalized Vessel-Level Input", styles["Heading2"]),
        ]
    )
    input_rows = [
        ["Formulation", "Vessel", "Time", "Percent released"],
        *[
            [
                str(row["formulation"]),
                str(row["batch"]),
                f"{float(row['time']):g}",
                f"{float(row['percent_released']):.3f}",
            ]
            for row in result.normalized_rows
        ],
    ]
    input_table = Table(input_rows, repeatRows=1)
    input_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dcebf7")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#aebdcc")),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.extend(
        [
            input_table,
            Spacer(1, 14),
            Paragraph("Exact Configuration", styles["Heading2"]),
            Paragraph(
                html.escape(json.dumps(result.config.to_dict(), sort_keys=True)),
                styles["Code"],
            ),
            Spacer(1, 14),
            Paragraph(f"<b>Disclaimer.</b> {html.escape(_DISCLAIMER)}", styles["Normal"]),
        ]
    )
    document = SimpleDocTemplate(str(output_path), pagesize=A4)
    document.build(story)
    return output_path.read_bytes()


def _write_docx(result: DissolutionWorkbenchResult, output_path: Path) -> bytes:
    from docx import Document
    from docx.shared import Inches

    document = Document()
    document.add_heading("Advanced Dissolution Workbench", level=0)
    document.add_paragraph(
        f"{result.config.reference_label} vs {result.config.test_label} | OpenPKFlow {__version__}"
    )
    summary = document.add_table(rows=1, cols=3)
    summary.style = "Table Grid"
    for index, value in enumerate(("Method", "Result", "Decision")):
        summary.rows[0].cells[index].text = value
    summary_rows = [
        (
            f"f2 ({result.comparison.f2_method})",
            f"{result.comparison.f2_value:.3f}",
            "Supports similarity"
            if result.comparison.f2_value >= 50.0
            else "Does not support similarity",
        ),
        (
            f"Bootstrap {result.bootstrap.confidence_level:.0%} CI",
            f"{result.bootstrap.ci_lower:.3f} - {result.bootstrap.ci_upper:.3f}",
            "Supports similarity" if result.bootstrap.is_similar else "Does not support similarity",
        ),
        (
            "MSD squared / critical",
            f"{result.msd_result.msd_squared:.3f} / {result.msd_result.chi2_05_critical:.3f}",
            "Supports similarity"
            if result.msd_result.is_similar
            else "Does not support similarity",
        ),
    ]
    for summary_values in summary_rows:
        cells = summary.add_row().cells
        for index, value in enumerate(summary_values):
            cells[index].text = value

    plot_path = output_path.with_suffix(".profile.png")
    plot_path.write_bytes(_profile_plot_png(result))
    document.add_heading("Mean and Vessel Profiles", level=1)
    document.add_picture(str(plot_path), width=Inches(6.4))
    plot_path.unlink(missing_ok=True)

    document.add_heading("Five-Model AICc Ranking", level=1)
    models = document.add_table(rows=1, cols=6)
    models.style = "Table Grid"
    for index, value in enumerate(("Formulation", "Rank", "Model", "Parameters", "R2", "AICc")):
        models.rows[0].cells[index].text = value
    for model_values in _model_rows(result):
        cells = models.add_row().cells
        for index, value in enumerate(model_values):
            cells[index].text = value

    document.add_heading("Normalized Vessel-Level Input", level=1)
    inputs = document.add_table(rows=1, cols=4)
    inputs.style = "Table Grid"
    for index, value in enumerate(("Formulation", "Vessel", "Time", "Percent released")):
        inputs.rows[0].cells[index].text = value
    for row in result.normalized_rows:
        input_values = (
            str(row["formulation"]),
            str(row["batch"]),
            f"{float(row['time']):g}",
            f"{float(row['percent_released']):.3f}",
        )
        cells = inputs.add_row().cells
        for index, value in enumerate(input_values):
            cells[index].text = value

    document.add_heading("Exact Configuration", level=1)
    document.add_paragraph(json.dumps(result.config.to_dict(), indent=2, sort_keys=True))
    document.add_heading("Disclaimer", level=1)
    document.add_paragraph(_DISCLAIMER)
    document.save(str(output_path))
    return output_path.read_bytes()


def report_workbench(
    result: DissolutionWorkbenchResult,
    output_path: str | Path,
    *,
    format: Literal["html", "pdf", "docx"] = "html",
) -> str | bytes:
    """Write a complete workbench report.

    Parameters
    ----------
    result : DissolutionWorkbenchResult
        Completed result.
    output_path : str | Path
        Destination path.
    format : {"html", "pdf", "docx"}, optional
        Output format.

    Returns
    -------
    str | bytes
        Rendered content.
    """
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    if format == "html":
        rendered = render_workbench_html(result)
        out.write_text(rendered, encoding="utf-8")
        return rendered
    if format == "pdf":
        return _write_pdf(result, out)
    if format == "docx":
        return _write_docx(result, out)
    raise ValueError("format must be 'html', 'pdf', or 'docx'.")


def _normalized_csv(result: DissolutionWorkbenchResult) -> bytes:
    stream = io.StringIO(newline="")
    writer = csv.DictWriter(
        stream,
        fieldnames=["formulation", "batch", "time", "percent_released"],
        lineterminator="\n",
    )
    writer.writeheader()
    writer.writerows(result.normalized_rows)
    return stream.getvalue().encode("utf-8")


def _writestr_deterministic(
    archive: zipfile.ZipFile,
    name: str,
    data: bytes,
) -> None:
    info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
    info.compress_type = zipfile.ZIP_DEFLATED
    info.external_attr = 0o600 << 16
    archive.writestr(info, data)


def write_workbench_audit_bundle(
    result: DissolutionWorkbenchResult,
    output_path: str | Path,
) -> Path:
    """Write normalized input, config, results, report, and SHA-256 manifest.

    Parameters
    ----------
    result : DissolutionWorkbenchResult
        Completed workbench result.
    output_path : str | Path
        Destination ZIP path.

    Returns
    -------
    Path
        Resolved archive path.
    """
    out = Path(output_path).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    files = {
        "input/normalized_dissolution.csv": _normalized_csv(result),
        "config.json": json.dumps(
            result.config.to_dict(),
            indent=2,
            sort_keys=True,
            allow_nan=False,
        ).encode("utf-8"),
        "results.json": workbench_result_json(result).encode("utf-8"),
        "report.html": render_workbench_html(result).encode("utf-8"),
    }
    manifest = {
        "openpkflow_version": __version__,
        "generated_at_utc": result.generated_at_utc,
        "hash_algorithm": "sha256",
        "files": {
            name: {
                "sha256": hashlib.sha256(data).hexdigest(),
                "size_bytes": len(data),
            }
            for name, data in sorted(files.items())
        },
    }
    files["manifest.json"] = json.dumps(
        manifest,
        indent=2,
        sort_keys=True,
        allow_nan=False,
    ).encode("utf-8")
    with zipfile.ZipFile(out, "w") as archive:
        for name, data in sorted(files.items()):
            _writestr_deterministic(archive, name, data)
    return out
