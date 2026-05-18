"""Tests for Word document report renderers.

Reference: openpkflow[reports] optional extra (python-docx >= 1.1).
DOCX magic bytes: b"PK\\x03\\x04"  (ZIP local file header, per ECMA-376)
"""
from __future__ import annotations

import pytest

docx = pytest.importorskip("docx", reason="python-docx not installed")

from openpkflow.report.docx import render_comparison_docx_report, render_model_fit_docx_report  # noqa: E402

_TIME = [5.0, 10.0, 15.0, 30.0, 45.0, 60.0]
_REF = [20.0, 40.0, 55.0, 70.0, 82.0, 90.0]
_TST = [18.0, 38.0, 53.0, 68.0, 80.0, 88.0]

_FIT_ROWS = [
    {
        "model_name": "weibull",
        "params": {"alpha": 1.5, "beta": 20.0},
        "r_squared": 0.998,
        "aic": -12.3,
        "aicc": -10.1,
        "bic": -11.5,
        "n_points": 6,
        "n_params": 2,
        "converged": True,
        "rank": 1,
        "is_best": True,
    },
    {
        "model_name": "first_order",
        "params": {"k1": 0.05},
        "r_squared": 0.985,
        "aic": -8.0,
        "aicc": -7.0,
        "bic": -7.8,
        "n_points": 6,
        "n_params": 1,
        "converged": True,
        "rank": 2,
        "is_best": False,
    },
    {
        "model_name": "zero_order",
        "params": {},
        "r_squared": float("nan"),
        "aic": float("nan"),
        "aicc": float("nan"),
        "bic": float("nan"),
        "n_points": 6,
        "n_params": 1,
        "converged": False,
        "rank": None,
        "is_best": False,
    },
]


class TestComparisonDOCX:
    def test_returns_bytes(self) -> None:
        result = render_comparison_docx_report(
            title="Test Report",
            reference_label="Ref",
            test_label="Test",
            f1_value=3.2,
            f2_value=68.5,
            n_timepoints=6,
            time_points=_TIME,
            reference_mean=_REF,
            test_mean=_TST,
        )
        assert isinstance(result, bytes)

    def test_docx_magic_bytes(self) -> None:
        result = render_comparison_docx_report(
            title="Test Report",
            reference_label="Ref",
            test_label="Test",
            f1_value=3.2,
            f2_value=68.5,
            n_timepoints=6,
            time_points=_TIME,
            reference_mean=_REF,
            test_mean=_TST,
        )
        assert result[:4] == b"PK\x03\x04"

    def test_minimum_size(self) -> None:
        result = render_comparison_docx_report(
            title="Test Report",
            reference_label="Ref",
            test_label="Test",
            f1_value=3.2,
            f2_value=68.5,
            n_timepoints=6,
            time_points=_TIME,
            reference_mean=_REF,
            test_mean=_TST,
        )
        assert len(result) > 1024

    def test_roundtrip_disclaimer(self, tmp_path: pytest.TempPathFactory) -> None:  # type: ignore[type-arg]
        import docx as _docx

        out = tmp_path / "comparison.docx"  # type: ignore[operator]
        render_comparison_docx_report(
            title="Test Report",
            reference_label="Ref",
            test_label="Test",
            f1_value=3.2,
            f2_value=68.5,
            n_timepoints=6,
            time_points=_TIME,
            reference_mean=_REF,
            test_mean=_TST,
            output_path=out,
        )
        assert out.exists()
        doc = _docx.Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "OpenPKFlow" in full_text
        assert "regulatory" in full_text.lower()

    def test_writes_file(self, tmp_path: pytest.TempPathFactory) -> None:  # type: ignore[type-arg]
        out = tmp_path / "comparison.docx"  # type: ignore[operator]
        render_comparison_docx_report(
            title="Test Report",
            reference_label="Ref",
            test_label="Test",
            f1_value=3.2,
            f2_value=68.5,
            n_timepoints=6,
            time_points=_TIME,
            reference_mean=_REF,
            test_mean=_TST,
            output_path=out,
        )
        assert out.exists()
        assert out.read_bytes()[:4] == b"PK\x03\x04"


class TestModelFitDOCX:
    def _make_plot_b64(self) -> str:
        import numpy as np

        from openpkflow.dissolution.plotting import dissolution_fit_plot_b64

        t_dense = list(np.linspace(0.0, 60.0, 50))
        fit_curves = [
            ("weibull", t_dense, [float(v) for v in np.linspace(0, 90, 50)], -10.1),
        ]
        return dissolution_fit_plot_b64(
            time_points=_TIME,
            observed_mean=_REF,
            fit_curves=fit_curves,
        )

    def test_returns_bytes(self) -> None:
        plot_b64 = self._make_plot_b64()
        result = render_model_fit_docx_report(
            formulation_label="Reference",
            time_points=_TIME,
            observed_mean=_REF,
            fit_rows=_FIT_ROWS,
            plot_b64=plot_b64,
        )
        assert isinstance(result, bytes)
        assert result[:4] == b"PK\x03\x04"

    def test_roundtrip_disclaimer(self, tmp_path: pytest.TempPathFactory) -> None:  # type: ignore[type-arg]
        import docx as _docx

        plot_b64 = self._make_plot_b64()
        out = tmp_path / "fit.docx"  # type: ignore[operator]
        render_model_fit_docx_report(
            formulation_label="Reference",
            time_points=_TIME,
            observed_mean=_REF,
            fit_rows=_FIT_ROWS,
            plot_b64=plot_b64,
            output_path=out,
        )
        assert out.exists()
        doc = _docx.Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "OpenPKFlow" in full_text
        assert "regulatory" in full_text.lower()
        assert "similarity test" in full_text.lower()

    def test_failed_models_noted(self) -> None:
        import io

        import docx as _docx

        plot_b64 = self._make_plot_b64()
        result = render_model_fit_docx_report(
            formulation_label="Reference",
            time_points=_TIME,
            observed_mean=_REF,
            fit_rows=_FIT_ROWS,
            plot_b64=plot_b64,
        )
        doc = _docx.Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "zero_order" in full_text

    def test_writes_file(self, tmp_path: pytest.TempPathFactory) -> None:  # type: ignore[type-arg]
        plot_b64 = self._make_plot_b64()
        out = tmp_path / "fit.docx"  # type: ignore[operator]
        render_model_fit_docx_report(
            formulation_label="Reference",
            time_points=_TIME,
            observed_mean=_REF,
            fit_rows=_FIT_ROWS,
            plot_b64=plot_b64,
            output_path=out,
        )
        assert out.exists()
        assert out.read_bytes()[:4] == b"PK\x03\x04"


class TestDOCXImportGuard:
    def test_import_guard_message(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import sys

        original = sys.modules.get("docx")
        sys.modules["docx"] = None  # type: ignore[assignment]
        try:
            import importlib

            import openpkflow.report.docx as docx_mod
            importlib.reload(docx_mod)
            with pytest.raises(ImportError, match="pip install openpkflow\\[reports\\]"):
                docx_mod.render_comparison_docx_report(
                    title="t",
                    reference_label="r",
                    test_label="t",
                    f1_value=0.0,
                    f2_value=50.0,
                    n_timepoints=1,
                    time_points=[5.0],
                    reference_mean=[50.0],
                    test_mean=[50.0],
                )
        finally:
            if original is None:
                sys.modules.pop("docx", None)
            else:
                sys.modules["docx"] = original
